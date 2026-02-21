from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BLUE = RGBColor(0x1A, 0x56, 0xDB)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x44, 0x44, 0x44)
GRAY_MED = RGBColor(0x55, 0x55, 0x55)
BODY_COLOR = RGBColor(0x33, 0x33, 0x33)

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(9.5)
style.font.color.rgb = DARK
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.15

for section in doc.sections:
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)
    header = section.header
    header.is_linked_to_previous = True
    for p in header.paragraphs:
        p.clear()
    footer = section.footer
    footer.is_linked_to_previous = True
    for p in footer.paragraphs:
        p.clear()


def add_section_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    border_el = p._p.get_or_add_pPr()
    pBdr = border_el.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "1",
            qn("w:color"): "CBD5E1",
        },
    )
    pBdr.append(bottom)
    border_el.append(pBdr)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE
    run.font.all_caps = True


def add_bullet(text, indent=Inches(0.25)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = indent
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(1)
    run_marker = p.add_run("\u25b8 ")
    run_marker.font.size = Pt(8)
    run_marker.font.color.rgb = BLUE
    run_text = p.add_run(text)
    run_text.font.size = Pt(9)
    run_text.font.color.rgb = BODY_COLOR


def add_job(title, dates, company):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(9.5)
    run_title.font.color.rgb = DARK
    p.add_run("\t")
    run_dates = p.add_run(dates)
    run_dates.font.size = Pt(8.5)
    run_dates.font.color.rgb = GRAY_MED
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.4), WD_ALIGN_PARAGRAPH.RIGHT)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    run_co = p2.add_run(company)
    run_co.italic = True
    run_co.font.size = Pt(9)
    run_co.font.color.rgb = GRAY_MED


# ===== HEADER =====
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.paragraph_format.space_after = Pt(2)
run_name = name.add_run("HARNEET K. PHUL")
run_name.bold = True
run_name.font.size = Pt(20)
run_name.font.color.rgb = BLUE

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(4)
border_el = contact._p.get_or_add_pPr()
pBdr = border_el.makeelement(qn("w:pBdr"), {})
bottom = pBdr.makeelement(
    qn("w:bottom"),
    {
        qn("w:val"): "single",
        qn("w:sz"): "12",
        qn("w:space"): "4",
        qn("w:color"): "1A56DB",
    },
)
pBdr.append(bottom)
border_el.append(pBdr)
run_contact = contact.add_run("917-207-2960  |  harneetkaur148@gmail.com")
run_contact.font.size = Pt(9)
run_contact.font.color.rgb = GRAY

# ===== PROFESSIONAL SUMMARY =====
add_section_title("Professional Summary")
summary = doc.add_paragraph()
summary.paragraph_format.space_after = Pt(2)
run_s = summary.add_run(
    "Software engineer with 14+ years of experience in test automation, data validation, "
    "and large-scale system integration, now pivoting into AI/ML engineering. Skilled in "
    "Python, TensorFlow, and NLP pipelines, with hands-on training in deep learning, "
    "generative AI, and prompt engineering. Strong foundation in building automated "
    "frameworks, analyzing complex datasets, and deploying quality-driven solutions in "
    "Agile environments. Azure AI Engineer Associate certified."
)
run_s.font.size = Pt(9)
run_s.font.color.rgb = BODY_COLOR

# ===== TECHNICAL SKILLS =====
add_section_title("Technical Skills")

skills = [
    ("AI / ML:", "TensorFlow, NLP, Deep Learning, Generative AI, Prompt Engineering, LLMs"),
    ("Languages:", "Python, Java, SQL, XSLT"),
    ("Cloud / DevOps:", "Azure AI Services, Google Cloud AI, Jenkins CI/CD"),
    ("Automation:", "Selenium WebDriver, BDD/Cucumber, REST API Testing"),
    ("Data / Tools:", "TOAD, SQL Databases, SOAP/REST, JIRA, Git"),
    ("Methodologies:", "Agile/Scrum, Test-Driven Development, CI/CD"),
]

table = doc.add_table(rows=3, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for row_idx in range(3):
    for col_idx in range(2):
        skill_idx = row_idx * 2 + col_idx
        cell = table.cell(row_idx, col_idx)
        cell.width = Inches(3.2)
        for p in cell.paragraphs:
            p.clear()
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        run_label = p.add_run(skills[skill_idx][0] + " ")
        run_label.bold = True
        run_label.font.size = Pt(9)
        run_label.font.color.rgb = BODY_COLOR
        run_val = p.add_run(skills[skill_idx][1])
        run_val.font.size = Pt(9)
        run_val.font.color.rgb = GRAY

for row in table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = tcPr.makeelement(qn("w:tcBorders"), {})
        for edge in ("top", "left", "bottom", "right"):
            el = borders.makeelement(
                qn(f"w:{edge}"), {qn("w:val"): "none", qn("w:sz"): "0", qn("w:space"): "0"}
            )
            borders.append(el)
        tcPr.append(borders)

# ===== CERTIFICATIONS =====
add_section_title("AI & Engineering Certifications")

certs = [
    "Azure AI Engineer Associate (AI-102)",
    "Machine Learning Specialization \u2013 DeepLearning.AI (Coursera)",
    "Google Cloud \u2013 Generative AI",
    "TensorFlow & NLP Fundamentals",
    "Generative AI Prompt Engineering",
    "Deep Learning Essentials",
    "Python Automation & Testing",
    "BDD \u2013 Cucumber Essentials",
]

cert_table = doc.add_table(rows=4, cols=2)
cert_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, cert in enumerate(certs):
    row_idx = i // 2
    col_idx = i % 2
    cell = cert_table.cell(row_idx, col_idx)
    cell.width = Inches(3.2)
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run_m = p.add_run("\u25b8 ")
    run_m.font.size = Pt(8)
    run_m.font.color.rgb = BLUE
    run_c = p.add_run(cert)
    run_c.font.size = Pt(9)
    run_c.font.color.rgb = BODY_COLOR

for row in cert_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = tcPr.makeelement(qn("w:tcBorders"), {})
        for edge in ("top", "left", "bottom", "right"):
            el = borders.makeelement(
                qn(f"w:{edge}"), {qn("w:val"): "none", qn("w:sz"): "0", qn("w:space"): "0"}
            )
            borders.append(el)
        tcPr.append(borders)

# ===== PROFESSIONAL EXPERIENCE =====
add_section_title("Professional Experience")

add_job("QA Consultant \u2013 AI & Automation Focus", "Oct 2018 \u2013 Present", "First Derivatives, Mississauga ON / New York, USA")
add_bullet("Designed and implemented automated test suites and data-driven frameworks using Python and Selenium, reducing manual test effort by 60%")
add_bullet("Built API test strategies for complex energy-sector systems (IESO), validating data integrity across scheduling, dispatch, and contract modules")
add_bullet("Authored automation scripts processing large datasets for regression and integration testing, establishing repeatable data pipelines")
add_bullet("Led Agile ceremonies and cross-functional collaboration, driving quality metrics analysis and continuous process improvement")

add_job("Sr. Software Quality Engineer", "Dec 2013 \u2013 Sep 2018", "Rogers Communications, Toronto, ON")
add_bullet("Engineered automated testing solutions with Selenium-Java and HP QTP, executing data validation across multi-tier e-commerce platforms")
add_bullet("Built test data repositories and SQL-based validation pipelines for system and integration testing across web applications")
add_bullet("Performed web service testing (SOAP/REST APIs) and root-cause analysis using log/trace file extraction and pattern recognition")
add_bullet("Led a team of 6 onsite/offshore engineers, coordinating technical assessments and effort estimation for delivery")

add_job("Software QA Engineer", "Jun 2013 \u2013 Dec 2013", "Xenex Enterprises Inc., Toronto, ON")
add_bullet("Developed complex automated test suites using Selenium WebDriver for web and mobile apps, including cross-browser compatibility")
add_bullet("Designed data-driven test scenarios for validation of business-critical workflows in integrated environments")

add_job("Jr. Software Engineer", "Jan 2011 \u2013 Feb 2013", "Information Mosaic / Miri Infotech, India")
add_bullet("Wrote SQL queries for backend data validation and created XSLT transformations for structured data processing")
add_bullet("Executed functional and regression testing, collaborating with development teams for rapid defect resolution")

# ===== EDUCATION =====
add_section_title("Education")
edu = doc.add_paragraph()
edu.paragraph_format.space_after = Pt(0)
edu.paragraph_format.tab_stops.add_tab_stop(Inches(6.4), WD_ALIGN_PARAGRAPH.RIGHT)
run_deg = edu.add_run("B.E. \u2013 Computer Science & Engineering")
run_deg.bold = True
run_deg.font.size = Pt(9)
run_deg.font.color.rgb = BODY_COLOR
edu.add_run("\t")
run_uni = edu.add_run("Punjab Technical University, 2006\u20132010")
run_uni.font.size = Pt(9)
run_uni.font.color.rgb = BODY_COLOR

doc.save("/workspace/resume.docx")
print("resume.docx created successfully")
