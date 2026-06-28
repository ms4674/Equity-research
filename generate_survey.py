#!/usr/bin/env python3
"""Generate a Word survey: 'LLM Token Consumption, Budget & Pricing Sensitivity Survey'.

Every question offers at least 4-5 answer options.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Palette / style helpers
# ---------------------------------------------------------------------------
NAVY = RGBColor(0x1F, 0x33, 0x5E)
ACCENT = RGBColor(0x2E, 0x6D, 0xB4)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT = RGBColor(0x8A, 0x8A, 0x8A)


def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    run = p.add_run(f"Section {number}.  {title}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "2E6DB4")
    pbdr.append(bottom)
    pPr.append(pbdr)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_intro_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(8)
    return p


def add_question(doc, qnum, text, note=None, multi=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"Q{qnum}. ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    run2 = p.add_run(text)
    run2.bold = True
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    if note:
        pn = doc.add_paragraph()
        rn = pn.add_run(note)
        rn.italic = True
        rn.font.size = Pt(9)
        rn.font.color.rgb = LIGHT
        pn.paragraph_format.space_after = Pt(2)
    elif multi:
        pn = doc.add_paragraph()
        rn = pn.add_run("(Select all that apply)")
        rn.italic = True
        rn.font.size = Pt(9)
        rn.font.color.rgb = LIGHT
        pn.paragraph_format.space_after = Pt(2)
    return p


def add_options(doc, options):
    for opt in options:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run("\u2610  ")  # ballot box
        run.font.size = Pt(11)
        run2 = p.add_run(opt)
        run2.font.size = Pt(10.5)
        run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_other_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("\u2610  Other (please specify): ____________________________________")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_matrix(doc, row_label, columns, rows):
    """Matrix-style table where each row is rated against column options (radio columns)."""
    table = doc.add_table(rows=1, cols=len(columns) + 1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = ""
    run = hdr[0].paragraphs[0].add_run(row_label)
    run.bold = True
    run.font.size = Pt(9)
    set_cell_background(hdr[0], "1F335E")
    hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, col in enumerate(columns, start=1):
        cell = hdr[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(col)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, "2E6DB4")
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        cells[0].text = ""
        r = cells[0].paragraphs[0].add_run(row)
        r.font.size = Pt(9.5)
        if ri % 2 == 1:
            set_cell_background(cells[0], "EEF2F8")
        for i in range(1, len(columns) + 1):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run("\u25CB")
            rr.font.size = Pt(11)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri % 2 == 1:
                set_cell_background(cells[i], "EEF2F8")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------
doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# ---- Title block ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
trun = title.add_run("LLM Token Consumption, Budget & Pricing Sensitivity Survey")
trun.bold = True
trun.font.size = Pt(20)
trun.font.color.rgb = NAVY
title.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph()
srun = sub.add_run("Benchmarking generative-AI token usage across use cases and sectors")
srun.italic = True
srun.font.size = Pt(12)
srun.font.color.rgb = ACCENT
sub.paragraph_format.space_after = Pt(10)

purpose = doc.add_paragraph()
prun = purpose.add_run(
    "Purpose:  This survey measures how organizations consume large language model (LLM) tokens "
    "by use case and by sector, your current token budget, and how sensitive your usage is to changes "
    "in token pricing. Responses are confidential and will be reported only in aggregate. "
    "Estimated completion time: 10\u201312 minutes."
)
prun.font.size = Pt(10.5)
prun.font.color.rgb = GREY
purpose.paragraph_format.space_after = Pt(4)

defn = doc.add_paragraph()
drun = defn.add_run(
    "Quick definition:  A \u201ctoken\u201d is the unit LLM providers use to price and meter text (and increasingly "
    "image/audio) input and output. Roughly 1,000 tokens \u2248 750 English words. \u201cInput\u201d (prompt) and "
    "\u201coutput\u201d (completion) tokens are usually priced separately."
)
drun.font.size = Pt(9.5)
drun.italic = True
drun.font.color.rgb = LIGHT

# =====================================================================
# SECTION A — Respondent & Organization Profile
# =====================================================================
add_section_heading(doc, "A", "Respondent & Organization Profile")
add_intro_text(doc, "Help us segment results by sector, organization size, and your role.")

add_question(doc, 1, "Which sector best describes your organization's primary industry?")
add_options(doc, [
    "Financial services, banking & insurance",
    "Technology, software & internet",
    "Healthcare, pharma & life sciences",
    "Retail, e-commerce & consumer goods",
    "Manufacturing, industrials & energy",
    "Media, marketing, entertainment & gaming",
    "Professional, legal & business services",
    "Public sector, education & non-profit",
    "Telecommunications & networking",
])
add_other_line(doc)

add_question(doc, 2, "What is the approximate size of your organization (full-time employees)?")
add_options(doc, [
    "1\u201350 (startup / small business)",
    "51\u2013250 (small-medium)",
    "251\u20131,000 (mid-market)",
    "1,001\u20135,000 (large enterprise)",
    "5,001\u201320,000 (large enterprise)",
    "More than 20,000 (global enterprise)",
])

add_question(doc, 3, "Which best describes your role relative to LLM/AI initiatives?")
add_options(doc, [
    "Executive sponsor / budget owner (CxO, VP)",
    "Engineering / data science / ML practitioner",
    "Product or program manager for AI features",
    "IT, platform, or cloud infrastructure / FinOps",
    "Procurement, finance, or vendor management",
    "Individual business user of AI tools",
])
add_other_line(doc)

add_question(doc, 4, "What is your organization's overall stage of LLM adoption?")
add_options(doc, [
    "Not yet using LLMs / exploring only",
    "Piloting one or a few proofs-of-concept",
    "In production for a single use case",
    "In production across multiple use cases",
    "Mature, organization-wide deployment with governance",
])

# =====================================================================
# SECTION B — Use Cases & Deployment
# =====================================================================
add_section_heading(doc, "B", "LLM Use Cases & Deployment Model")
add_intro_text(doc, "Tell us where and how you apply LLMs today.")

add_question(doc, 5, "Which LLM use cases are active in your organization today?", multi=True)
add_options(doc, [
    "Customer support / chatbots / virtual agents",
    "Software development & coding assistants",
    "Content creation, marketing & copywriting",
    "Search, retrieval-augmented generation (RAG) & knowledge management",
    "Data analysis, BI, summarization & reporting",
    "Document processing, extraction & classification",
    "Autonomous / agentic workflows & task automation",
    "Translation, transcription & multilingual support",
    "Internal productivity copilots (email, meetings, docs)",
])
add_other_line(doc)

add_question(doc, 6, "Which single use case consumes the MOST tokens in your organization?")
add_options(doc, [
    "Customer support / conversational agents",
    "Coding & software development assistants",
    "Content & marketing generation",
    "RAG / search / knowledge retrieval",
    "Document & data processing at scale",
    "Agentic / multi-step automation workflows",
])
add_other_line(doc)

add_question(doc, 7, "How do you primarily access LLM capabilities?")
add_options(doc, [
    "Commercial API providers (e.g., OpenAI, Anthropic, Google)",
    "Cloud-hosted models (Azure, AWS Bedrock, Vertex AI)",
    "Self-hosted open-weight models (Llama, Mistral, Qwen, etc.)",
    "Packaged SaaS products with embedded LLMs",
    "A mix of API, self-hosted, and SaaS approaches",
])
add_other_line(doc)

add_question(doc, 8, "How many distinct LLM-powered applications/workflows are in production?")
add_options(doc, [
    "None yet (pilots only)",
    "1\u20132",
    "3\u20135",
    "6\u201310",
    "11\u201325",
    "More than 25",
])

add_question(doc, 9, "What is the balance between open-source/open-weight and proprietary (closed) models in your token usage?")
add_options(doc, [
    "Exclusively proprietary / closed models (e.g., GPT, Claude, Gemini)",
    "Mostly proprietary, some open-source (\u2265 60% proprietary)",
    "Roughly balanced mix of open-source and proprietary",
    "Mostly open-source / open-weight, some proprietary (\u2265 60% open)",
    "Exclusively open-source / open-weight models (e.g., Llama, Mistral, Qwen)",
    "Don't know / not tracked",
])

add_question(doc, 10, "To what extent do you use non-US / Chinese-origin models (e.g., DeepSeek, Kimi, GLM, Qwen) versus US models (e.g., GPT, Claude, Gemini, Llama)?")
add_options(doc, [
    "Exclusively US models",
    "Mostly US models, occasional testing of non-US models",
    "Roughly balanced use of US and non-US models",
    "Mostly non-US / Chinese-origin models",
    "Exclusively non-US / Chinese-origin models",
    "Evaluating non-US models but none in production yet",
])
add_other_line(doc)

# =====================================================================
# SECTION C — Token Consumption by Use Case (matrix)
# =====================================================================
add_section_heading(doc, "C", "Token Consumption by Use Case")
add_intro_text(doc, "For each use case you operate, estimate the monthly token volume it consumes. "
                    "Mark one circle per row; leave a row blank if not applicable.")

add_question(doc, 11, "Estimated MONTHLY token consumption by use case:",
             note="Bands are total tokens/month (input + output). 1M = one million, 1B = one billion.")
vol_cols = ["< 1M", "1M\u201310M", "10M\u2013100M", "100M\u20131B", "> 1B", "N/A"]
add_matrix(doc, "Use case",
           vol_cols,
           [
               "Customer support / chatbots",
               "Coding assistants",
               "Content & marketing",
               "RAG / search / knowledge",
               "Data analysis & summarization",
               "Document processing & extraction",
               "Agentic / automation workflows",
           ])

add_question(doc, 12, "What is your organization's approximate TOTAL token consumption per month (all use cases combined)?")
add_options(doc, [
    "Less than 1 million tokens / month",
    "1 million \u2013 50 million tokens / month",
    "50 million \u2013 500 million tokens / month",
    "500 million \u2013 5 billion tokens / month",
    "5 billion \u2013 50 billion tokens / month",
    "More than 50 billion tokens / month",
    "Not measured / don't know",
])

add_question(doc, 13, "What is the approximate split between input (prompt) and output (completion) tokens?")
add_options(doc, [
    "Mostly input (\u2265 80% input)",
    "Input-heavy (roughly 60\u201380% input)",
    "Balanced (roughly 40\u201360% input)",
    "Output-heavy (roughly 60\u201380% output)",
    "Mostly output (\u2265 80% output)",
    "Don't know / not tracked",
])

add_question(doc, 14, "How has your total token consumption changed over the past 12 months?")
add_options(doc, [
    "Decreased",
    "Flat (roughly unchanged)",
    "Grew up to 2\u00d7",
    "Grew 2\u00d7 to 5\u00d7",
    "Grew 5\u00d7 to 10\u00d7",
    "Grew more than 10\u00d7",
    "Not applicable / too new to compare",
])

# =====================================================================
# SECTION D — Token Budget
# =====================================================================
add_section_heading(doc, "D", "Current Token / LLM Budget")
add_intro_text(doc, "Help us understand the financial scale of your LLM usage.")

add_question(doc, 15, "What is your organization's approximate ANNUAL spend on LLM tokens / API usage (USD)?")
add_options(doc, [
    "Less than $10,000",
    "$10,000 \u2013 $100,000",
    "$100,000 \u2013 $500,000",
    "$500,000 \u2013 $2 million",
    "$2 million \u2013 $10 million",
    "More than $10 million",
    "Don't know / not disclosed",
])

add_question(doc, 16, "How is your LLM/token budget primarily managed?")
add_options(doc, [
    "Centralized platform / AI team budget",
    "Allocated per business unit or product team",
    "Pooled within the cloud / infrastructure budget",
    "Ad hoc / no formal budget yet",
    "Managed by a dedicated FinOps / AI cost function",
])
add_other_line(doc)

add_question(doc, 17, "What share of your total IT / cloud budget does LLM token spend currently represent?")
add_options(doc, [
    "Less than 1%",
    "1% \u2013 5%",
    "5% \u2013 10%",
    "10% \u2013 25%",
    "More than 25%",
    "Don't know",
])

add_question(doc, 18, "How do you expect your LLM token budget to change over the NEXT 12 months?")
add_options(doc, [
    "Decrease",
    "Stay roughly flat",
    "Increase up to 25%",
    "Increase 25% \u2013 100%",
    "Increase 2\u00d7 \u2013 5\u00d7",
    "Increase more than 5\u00d7",
])

add_question(doc, 19, "Which methods do you actively use to control or reduce token costs?", multi=True)
add_options(doc, [
    "Prompt optimization / shorter prompts",
    "Caching (prompt/response caching)",
    "Routing to smaller / cheaper models where possible",
    "Self-hosting open-weight models",
    "Batching, rate limiting, or usage quotas",
    "Fine-tuning / distillation to reduce per-call tokens",
    "No active cost-control measures yet",
])
add_other_line(doc)

# =====================================================================
# SECTION E — Sensitivity to Token Pricing
# =====================================================================
add_section_heading(doc, "E", "Sensitivity to Token Pricing")
add_intro_text(doc, "Help us gauge how price changes would affect your usage and vendor choices.")

add_question(doc, 20, "How sensitive is your LLM usage to changes in token pricing?")
add_options(doc, [
    "Not sensitive \u2014 price is not a factor in our decisions",
    "Slightly sensitive \u2014 minor consideration",
    "Moderately sensitive \u2014 we monitor and react to changes",
    "Highly sensitive \u2014 price strongly drives our choices",
    "Critically sensitive \u2014 price is the primary constraint",
])

add_question(doc, 21, "If token prices INCREASED by 50%, what is the most likely response?")
add_options(doc, [
    "No change \u2014 absorb the cost",
    "Optimize prompts/caching to offset the increase",
    "Switch some workloads to cheaper / smaller models",
    "Switch to a different provider or self-hosting",
    "Reduce overall usage / pause some use cases",
    "Cancel or significantly scale back LLM initiatives",
])

add_question(doc, 22, "If token prices DECREASED by 50%, how would your usage most likely respond?")
add_options(doc, [
    "No meaningful change in usage",
    "Modest increase (up to 25% more usage)",
    "Significant increase (25% \u2013 100% more usage)",
    "Usage would more than double (2\u00d7+)",
    "We would launch new use cases previously deemed too costly",
])

add_question(doc, 23, "Which factors matter MOST when choosing an LLM provider/model?", multi=True)
add_options(doc, [
    "Token price / cost per request",
    "Model quality & accuracy",
    "Latency / response speed",
    "Context window size",
    "Data privacy, security & compliance",
    "Ease of integration & tooling / ecosystem",
    "Reliability, uptime & rate limits",
])
add_other_line(doc)

add_question(doc, 24, "What is the maximum acceptable price (per 1M output tokens) for your primary use case?")
add_options(doc, [
    "Under $1 per 1M tokens",
    "$1 \u2013 $5 per 1M tokens",
    "$5 \u2013 $15 per 1M tokens",
    "$15 \u2013 $50 per 1M tokens",
    "Over $50 per 1M tokens",
    "Price is not a binding constraint",
])

add_question(doc, 25, "How frequently do you re-evaluate model/provider choices based on price?")
add_options(doc, [
    "Rarely / never",
    "Annually",
    "Quarterly",
    "Monthly",
    "Continuously / automated routing by cost",
])

# =====================================================================
# SECTION F — Outlook & Open Feedback
# =====================================================================
add_section_heading(doc, "F", "Outlook & Open Feedback")

add_question(doc, 26, "What is the single biggest barrier to scaling your LLM usage?")
add_options(doc, [
    "Token / usage cost",
    "Model accuracy or reliability",
    "Data privacy, security & compliance",
    "Integration & engineering effort",
    "Lack of internal skills / talent",
    "Unclear ROI / business case",
])
add_other_line(doc)

add_question(doc, 27, "Over the next 2\u20133 years, what do you expect for token prices in the market?")
add_options(doc, [
    "Fall significantly (more than 50%)",
    "Fall moderately (10\u201350%)",
    "Stay roughly flat",
    "Rise moderately (10\u201350%)",
    "Rise significantly (more than 50%)",
    "No clear expectation",
])

# Open comment box
add_question(doc, 28, "Any additional comments on your token consumption, budget, or pricing sensitivity? (optional)")
for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BBBBBB")
    pbdr.append(bottom)
    pPr.append(pbdr)

# ---- Closing ----
close = doc.add_paragraph()
close.paragraph_format.space_before = Pt(10)
crun = close.add_run("Thank you for completing this survey. Your responses help benchmark LLM token "
                     "economics across the industry and will be shared only in aggregate, anonymized form.")
crun.italic = True
crun.font.size = Pt(10)
crun.font.color.rgb = GREY

doc.save("LLM_Token_Consumption_Survey.docx")
print("Saved LLM_Token_Consumption_Survey.docx")
